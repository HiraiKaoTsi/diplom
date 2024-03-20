from docx import Document
from docx.shared import Pt


def CreateReportMainInfo(way_and_name_file: str, info_meaning: dict) -> bool:
    """
    Создает doc файл, с информацией основного окна
    :param way_and_name_file: путь до файла и его название
    :param info_meaning: значение по которым будет заполняться основная информация
    """

    doc = Document()

    style = doc.styles['Normal']
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)

    head = doc.add_paragraph("Общая информация")
    head.runs[0].font.bold = True 
    head.runs[0].font.size = Pt(20)
    head.alignment = 1

    under_head = doc.add_paragraph("за 15.01.2023")
    under_head.alignment = 1

    doc.add_paragraph('')

    table = doc.add_table(rows=len(info_meaning)+2, cols=2)

    cell = table.cell(0, 0)
    cell.text = "Категории"
    cell.paragraphs[0].runs[0].font.bold = True 
    cell.paragraphs[0].runs[0].font.size = Pt(20) 
    cell.paragraphs[0].alignment = 1 

    cell = table.cell(0, 1) 
    cell.text = "Значение"
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(20)  
    cell.paragraphs[0].alignment = 1 

    row = 1
    for key in info_meaning:
        cell = table.cell(row, 0) 
        cell.text = f"{key}"

        cell = table.cell(row, 1) 
        cell.text = f"{info_meaning[key]}"

        row += 1

    doc.save(f"{way_and_name_file}")

    return True
